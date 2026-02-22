#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Tuple, Optional

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# -----------------------------
# Paths / Run bundle
# -----------------------------

def project_root() -> Path:
    # .../85_newwave_pipeline/scripts/run_paper_gen.py -> root = parents[1]
    return Path(__file__).resolve().parents[1]


@dataclass
class RunBundle:
    run_dir: Path
    methods_md: Path
    results_md: Path
    paper_md: Path
    paper_docx: Path


def new_run_dir(root: Path) -> Path:
    out = root / "outputs" / "prism_runs"
    out.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    rd = out / ts
    rd.mkdir(parents=True, exist_ok=True)
    return rd


# -----------------------------
# Font forcing (fix garbling)
# -----------------------------

def set_style_font(style, ascii_font: str, eastasia_font: str, size_pt: Optional[int] = None, bold: Optional[bool] = None):
    """Force style font for Word and for East Asia (Chinese) rendering."""
    style.font.name = ascii_font
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), eastasia_font)

    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold


def force_document_fonts(doc: Document, ascii_font: str = "Times New Roman", eastasia_font: str = "Songti SC"):
    """Apply robust font settings to common styles to avoid symbol-font garbling."""
    styles = doc.styles

    # Normal / Body
    for sname in ["Normal", "Body Text"]:
        if sname in [s.name for s in styles]:
            set_style_font(styles[sname], ascii_font, eastasia_font, size_pt=12, bold=False)

    # Headings
    for sname, sz in [("Title", 28), ("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        if sname in [s.name for s in styles]:
            set_style_font(styles[sname], ascii_font, eastasia_font, size_pt=sz, bold=True)

    # List styles (best effort)
    for sname in ["List Paragraph"]:
        if sname in [s.name for s in styles]:
            set_style_font(styles[sname], ascii_font, eastasia_font, size_pt=12, bold=False)


def apply_run_fonts(run, ascii_font: str = "Times New Roman", eastasia_font: str = "Songti SC", size_pt: Optional[int] = None, bold: Optional[bool] = None):
    run.font.name = ascii_font
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), eastasia_font)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


# -----------------------------
# Minimal Markdown -> DOCX
# (enough for methods/results drafts)
# -----------------------------

MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
MD_BULLET = re.compile(r"^(\s*)[-*]\s+(.*)$")
MD_NUMBERED = re.compile(r"^(\s*)\d+\.\s+(.*)$")


def add_md_block(doc: Document, md_text: str, ascii_font: str = "Times New Roman", eastasia_font: str = "Songti SC"):
    lines = md_text.replace("\r\n", "\n").split("\n")
    in_code = False
    code_buf: List[str] = []

    for raw in lines:
        line = raw.rstrip("\n")

        # fenced code
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                # flush code
                p = doc.add_paragraph(style="Normal")
                run = p.add_run("\n".join(code_buf))
                apply_run_fonts(run, ascii_font, eastasia_font, size_pt=10)
                in_code = False
                code_buf = []
            continue

        if in_code:
            code_buf.append(line)
            continue

        if not line.strip():
            doc.add_paragraph("")  # blank line
            continue

        # headings
        m = MD_HEADING.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level <= 1:
                p = doc.add_paragraph(text, style="Heading 1")
            elif level == 2:
                p = doc.add_paragraph(text, style="Heading 2")
            else:
                p = doc.add_paragraph(text, style="Heading 3")
            # ensure run fonts
            for run in p.runs:
                apply_run_fonts(run, ascii_font, eastasia_font)
            continue

        # bullet list
        mb = MD_BULLET.match(line)
        if mb:
            text = mb.group(2).strip()
            p = doc.add_paragraph(style="List Paragraph")
            run = p.add_run("• " + text)
            apply_run_fonts(run, ascii_font, eastasia_font)
            continue

        # numbered list
        mn = MD_NUMBERED.match(line)
        if mn:
            text = mn.group(2).strip()
            p = doc.add_paragraph(style="List Paragraph")
            run = p.add_run(text)
            apply_run_fonts(run, ascii_font, eastasia_font)
            continue

        # normal paragraph
        p = doc.add_paragraph(style="Normal")
        run = p.add_run(line)
        apply_run_fonts(run, ascii_font, eastasia_font)


# -----------------------------
# Template / defaults
# -----------------------------

DEFAULT_METHODS = """# Methods

## Data and scope
We compile a small event corpus related to the “85 New Wave” narrative and its later archival re-inscriptions. Each row is an event-like record (exhibition / publication / platform item), enriched with 4D iconology features.

## Feature construction (4D)
We operationalize four iconological directions:
- Time: rupture-related scores and component means.
- Space: city concentration and spatial scope / evidence.
- Social: institution-level categories (market / media / museum / platform / unknown).
- Tech/Channel: dissemination channels and their shares.

## O/R/S stage projection
We project the 4D features into a 3-state stage space:
- O (Openness): strength of objectification / archivalization.
- R (Reorganization): rupture / polemic intensity (plus semantic noise).
- S (SteadyState / opacity): structural opacity driven by missingness / unknown evidence.

We aggregate yearly means and normalized shares, and visualize trajectories across observed years.
"""

DEFAULT_RESULTS = """# Results

## Rupture components
We report yearly means for O/R/S and semantic noise, as well as yearly shares normalized within (O+R+S).

## Key figures
- fig_rupture_components_means.png
- fig_rupture_components_shares.png

## Tables
- table_rupture_components_yearly.csv
"""

DEFAULT_PAPER_HEADER = """# Paper Draft (auto-assembled)

This document is auto-assembled from the pipeline outputs. Replace this front matter with your abstract/introduction when ready.
"""


# -----------------------------
# Collect artifacts
# -----------------------------

def copy_if_exists(src: Path, dst_dir: Path) -> Optional[Path]:
    if not src.exists():
        return None
    dst = dst_dir / src.name
    shutil.copy2(src, dst)
    return dst


def collect_outputs(root: Path, run_dir: Path) -> List[Path]:
    """
    Copy a curated set of artifacts into the run folder.
    We include rupture outputs and also any 4d figures/tables if present.
    """
    outputs = root / "outputs"
    copied: List[Path] = []

    # Core rupture artifacts
    for name in [
        "fig_rupture_components_means.png",
        "fig_rupture_components_shares.png",
        "table_rupture_components_yearly.csv",
    ]:
        p = copy_if_exists(outputs / name, run_dir)
        if p:
            copied.append(p)

    # Also include other useful figures/tables if they exist
    patterns = [
        "fig_4d_*.png",
        "table_4d_*.csv",
        "table_time_rupture_yearly.csv",
        "table_geo_*audit.csv",
        "fig_polemic_vs_archival_threshold7.png",
        "table_yearly_polemic_vs_archival_threshold7.csv",
        "table_yearly_strong_archival_top_refs_threshold7.csv",
    ]
    for pat in patterns:
        for src in outputs.glob(pat):
            # Avoid duplicating core rupture artifacts
            if src.name in {p.name for p in copied}:
                continue
            p = copy_if_exists(src, run_dir)
            if p:
                copied.append(p)

    return copied


# -----------------------------
# Assemble MD files
# -----------------------------

def ensure_templates(run_dir: Path) -> Tuple[Path, Path, Path]:
    methods = run_dir / "methods.md"
    results = run_dir / "results.md"
    paper = run_dir / "paper.md"

    if not methods.exists():
        methods.write_text(DEFAULT_METHODS, encoding="utf-8")
    if not results.exists():
        results.write_text(DEFAULT_RESULTS, encoding="utf-8")
    if not paper.exists():
        # paper.md is concatenation
        paper.write_text(DEFAULT_PAPER_HEADER + "\n\n" + methods.read_text(encoding="utf-8") + "\n\n" + results.read_text(encoding="utf-8"),
                         encoding="utf-8")

    return methods, results, paper


# -----------------------------
# DOCX builder
# -----------------------------

def add_figure(doc: Document, fig_path: Path, caption: str, ascii_font: str, eastasia_font: str):
    if not fig_path.exists():
        return
    # Add image
    p = doc.add_paragraph("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    apply_run_fonts(run, ascii_font, eastasia_font)
    # width heuristic
    run.add_picture(str(fig_path), width=Inches(6.5))

    # Caption
    cap = doc.add_paragraph(style="Normal")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    apply_run_fonts(r, ascii_font, eastasia_font, size_pt=10, bold=False)


def add_table_preview(doc: Document, csv_path: Path, title: str, ascii_font: str, eastasia_font: str, max_rows: int = 12):
    if not csv_path.exists():
        return
    df = pd.read_csv(csv_path)
    doc.add_paragraph(title, style="Heading 3")

    # limit rows
    df2 = df.head(max_rows).copy()

    table = doc.add_table(rows=1, cols=len(df2.columns))
    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df2.columns):
        hdr_cells[j].text = str(col)
        # force font in header cell runs
        for p in hdr_cells[j].paragraphs:
            for r in p.runs:
                apply_run_fonts(r, ascii_font, eastasia_font, bold=True)

    for _, row in df2.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df2.columns):
            cells[j].text = str(row[col])
            for p in cells[j].paragraphs:
                for r in p.runs:
                    apply_run_fonts(r, ascii_font, eastasia_font)

    doc.add_paragraph(f"(Preview first {len(df2)} rows; full CSV saved alongside.)", style="Normal")
    # enforce font for that note
    for r in doc.paragraphs[-1].runs:
        apply_run_fonts(r, ascii_font, eastasia_font, size_pt=10)


def build_docx(bundle: RunBundle, copied: List[Path], ascii_font: str = "Times New Roman", eastasia_font: str = "Songti SC"):
    doc = Document()
    force_document_fonts(doc, ascii_font=ascii_font, eastasia_font=eastasia_font)

    # Title
    title = doc.add_paragraph("Paper Draft (auto-assembled)", style="Title")
    for r in title.runs:
        apply_run_fonts(r, ascii_font, eastasia_font)

    doc.add_paragraph("")

    # Methods / Results as markdown blocks
    methods_txt = bundle.methods_md.read_text(encoding="utf-8")
    results_txt = bundle.results_md.read_text(encoding="utf-8")

    doc.add_paragraph("Methods", style="Heading 1")
    add_md_block(doc, methods_txt, ascii_font=ascii_font, eastasia_font=eastasia_font)

    doc.add_paragraph("Results", style="Heading 1")
    add_md_block(doc, results_txt, ascii_font=ascii_font, eastasia_font=eastasia_font)

    # Figures section
    doc.add_paragraph("Figures and Tables", style="Heading 1")

    # Prefer rupture figures if present in run_dir
    fig_means = bundle.run_dir / "fig_rupture_components_means.png"
    fig_shares = bundle.run_dir / "fig_rupture_components_shares.png"
    tab_yearly = bundle.run_dir / "table_rupture_components_yearly.csv"

    if fig_means.exists():
        add_figure(doc, fig_means, "Figure 1. Rupture components (means) by year.", ascii_font, eastasia_font)
    if fig_shares.exists():
        add_figure(doc, fig_shares, "Figure 2. Rupture components (shares) by year.", ascii_font, eastasia_font)

    if tab_yearly.exists():
        add_table_preview(doc, tab_yearly, "Table 1. Rupture components yearly (preview)", ascii_font, eastasia_font)

    # Save
    doc.save(bundle.paper_docx)


# -----------------------------
# Main
# -----------------------------

def main():
    root = project_root()
    run_dir = new_run_dir(root)

    # Copy artifacts
    copied = collect_outputs(root, run_dir)

    # Ensure templates
    methods_md, results_md, paper_md = ensure_templates(run_dir)

    # Bundle object
    bundle = RunBundle(
        run_dir=run_dir,
        methods_md=methods_md,
        results_md=results_md,
        paper_md=paper_md,
        paper_docx=run_dir / "paper.docx",
    )

    # Build docx (try Songti SC; if user doesn't have it, Word will fall back;
    # you can change eastasia_font to SimSun if you prefer.)
    build_docx(bundle, copied, ascii_font="Times New Roman", eastasia_font="Songti SC")

    print("Wrote DOCX:", bundle.paper_docx)
    print("Done. Wrote paper bundle to:")
    print(" -", bundle.run_dir)
    print("Files:")
    print(" -", bundle.methods_md)
    print(" -", bundle.results_md)
    print(" -", bundle.paper_md)
    # list copied items (short)
    for p in copied[:12]:
        print(" -", p)
    if len(copied) > 12:
        print(f" - ... (+{len(copied)-12} more artifacts)")


if __name__ == "__main__":
    main()